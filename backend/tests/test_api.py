from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.config import settings
from app.auth.jwt_handler import create_access_token
from app.auth.password_handler import hash_password
from app.database.base import Base
from app.database.connection import get_db
from app.main import app
from app.models.password_reset_model import PasswordResetToken
from app.auth.reset_token_handler import hash_reset_token
from app.models.job_description_model import ActivityLog
from app.models.user_model import User


@pytest.fixture()
def client(tmp_path: Path):
    database_url = f"sqlite:///{tmp_path / 'test.db'}"
    engine = create_engine(database_url, connect_args={"check_same_thread": False})
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    Base.metadata.create_all(bind=engine)
    settings.upload_dir = tmp_path / "uploads"

    def override_get_db():
        db = TestingSessionLocal()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = override_get_db
    app.state.TestingSessionLocal = TestingSessionLocal
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


def register_and_login(client: TestClient) -> dict[str, str]:
    register_payload = {"name": "Test User", "email": "test@example.com", "password": "Password123"}
    response = client.post("/api/auth/register", json=register_payload)
    assert response.status_code == 201

    response = client.post("/api/auth/login", json={"email": "test@example.com", "password": "Password123"})
    assert response.status_code == 200
    token = response.json()["data"]["access_token"]
    return {"Authorization": f"Bearer {token}"}


def create_admin_headers(client: TestClient) -> dict[str, str]:
    db = client.app.state.TestingSessionLocal()
    try:
        admin = User(
            name="Admin User",
            email="admin@example.com",
            password_hash=hash_password("AdminPassword123"),
            role="admin",
        )
        db.add(admin)
        db.commit()
        db.refresh(admin)
        token = create_access_token(str(admin.id))
        return {"Authorization": f"Bearer {token}"}
    finally:
        db.close()


def make_docx_resume(path: Path) -> None:
    document = Document()
    document.add_heading("Contact", level=1)
    document.add_paragraph("test@example.com +1 555 123 4567 GitHub LinkedIn")
    document.add_heading("Summary", level=1)
    document.add_paragraph("Backend developer with practical FastAPI and Python project experience.")
    document.add_heading("Education", level=1)
    document.add_paragraph("BS Computer Science, Example University")
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python FastAPI SQLite Git GitHub Docker Communication Problem Solving")
    document.add_heading("Projects", level=1)
    document.add_paragraph(
        "Built an AI Resume Analyzer using FastAPI, SQLAlchemy, SQLite, PDF parsing, and ATS scoring."
    )
    document.save(path)


def make_pdf_resume(path: Path) -> None:
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 760, "Contact")
    pdf.drawString(72, 740, "test@example.com +1 555 123 4567 GitHub LinkedIn")
    pdf.drawString(72, 710, "Education")
    pdf.drawString(72, 690, "BS Computer Science, Example University")
    pdf.drawString(72, 660, "Skills")
    pdf.drawString(72, 640, "Python FastAPI SQLite Git GitHub Docker Communication")
    pdf.drawString(72, 610, "Projects")
    pdf.drawString(72, 590, "Built an AI Resume Analyzer with ATS scoring and job matching.")
    pdf.save()


def upload_valid_resume(client: TestClient, headers: dict[str, str], tmp_path: Path) -> int:
    resume_path = tmp_path / "resume.docx"
    make_docx_resume(resume_path)
    with resume_path.open("rb") as resume_file:
        response = client.post(
            "/api/resume/upload",
            headers=headers,
            files={
                "file": (
                    "resume.docx",
                    resume_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert response.status_code == 201
    return response.json()["data"]["id"]


def test_register(client: TestClient):
    response = client.post(
        "/api/auth/register",
        json={"name": "Test User", "email": "test@example.com", "password": "Password123"},
    )
    assert response.status_code == 201
    assert response.json()["success"] is True

    duplicate = client.post(
        "/api/auth/register",
        json={"name": "Test User", "email": "test@example.com", "password": "Password123"},
    )
    assert duplicate.status_code == 409
    assert duplicate.json()["message"] == "Email already registered. Please login or reset your password."


def test_login(client: TestClient):
    headers = register_and_login(client)
    assert headers["Authorization"].startswith("Bearer ")

    response = client.get("/api/auth/me", headers=headers)
    assert response.status_code == 200
    assert response.json()["data"]["email"] == "test@example.com"

    response = client.post("/api/auth/login", json={"email": "test@example.com", "password": "WrongPassword"})
    assert response.status_code == 401
    assert response.json()["message"] == "Invalid email or password."


def test_forgot_and_reset_password(client: TestClient):
    register_payload = {"name": "Reset User", "email": "reset@example.com", "password": "Password123"}
    response = client.post("/api/auth/register", json=register_payload)
    assert response.status_code == 201

    response = client.post("/api/auth/forgot-password", json={"email": "reset@example.com"})
    assert response.status_code == 200
    assert response.json()["message"] == "If this email exists, a password reset link has been sent."

    response = client.post("/api/auth/forgot-password", json={"email": "missing@example.com"})
    assert response.status_code == 200
    assert response.json()["message"] == "If this email exists, a password reset link has been sent."

    db = client.app.state.TestingSessionLocal()
    try:
        from app.models.user_model import User

        user = db.query(User).filter(User.email == "reset@example.com").first()
        reset_token = PasswordResetToken(
            user_id=user.id,
            token_hash=hash_reset_token("known-reset-token-for-test-value-123456"),
            expires_at=__import__("datetime").datetime.utcnow() + __import__("datetime").timedelta(minutes=30),
        )
        db.add(reset_token)
        db.commit()
    finally:
        db.close()

    response = client.post(
        "/api/auth/reset-password",
        json={"token": "known-reset-token-for-test-value-123456", "new_password": "NewPassword123"},
    )
    assert response.status_code == 200
    assert response.json()["message"] == "Password reset successfully. Please login."

    response = client.post("/api/auth/login", json={"email": "reset@example.com", "password": "NewPassword123"})
    assert response.status_code == 200


def test_upload_invalid_file(client: TestClient):
    headers = register_and_login(client)
    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("resume.txt", b"not a resume", "text/plain")},
    )
    assert response.status_code == 400
    assert response.json()["success"] is False


def test_upload_empty_pdf(client: TestClient):
    headers = register_and_login(client)
    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("empty.pdf", b"", "application/pdf")},
    )
    assert response.status_code == 400
    assert response.json()["success"] is False


def test_upload_corrupt_pdf(client: TestClient):
    headers = register_and_login(client)
    response = client.post(
        "/api/resume/upload",
        headers=headers,
        files={"file": ("broken.pdf", b"not a real pdf", "application/pdf")},
    )
    assert response.status_code == 422
    assert response.json()["success"] is False


def test_upload_oversized_file(client: TestClient):
    headers = register_and_login(client)
    original_limit = settings.max_upload_size_mb
    settings.max_upload_size_mb = 0
    try:
        response = client.post(
            "/api/resume/upload",
            headers=headers,
            files={
                "file": (
                    "resume.docx",
                    b"large enough for this zero megabyte test limit",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    finally:
        settings.max_upload_size_mb = original_limit
    assert response.status_code == 413
    assert response.json()["success"] is False


def test_upload_valid_resume(client: TestClient, tmp_path: Path):
    headers = register_and_login(client)
    resume_id = upload_valid_resume(client, headers, tmp_path)
    assert resume_id > 0


def test_upload_valid_pdf_resume(client: TestClient, tmp_path: Path):
    headers = register_and_login(client)
    resume_path = tmp_path / "resume.pdf"
    make_pdf_resume(resume_path)
    with resume_path.open("rb") as resume_file:
        response = client.post(
            "/api/resume/upload",
            headers=headers,
            files={"file": ("resume.pdf", resume_file, "application/pdf")},
        )
    assert response.status_code == 201
    assert response.json()["data"]["file_type"] == "pdf"


def test_analyze_resume_and_get_report(client: TestClient, tmp_path: Path):
    headers = register_and_login(client)
    resume_id = upload_valid_resume(client, headers, tmp_path)
    response = client.post(f"/api/analysis/resume/{resume_id}", headers=headers)
    assert response.status_code == 200
    report_id = response.json()["data"]["report_id"]

    response = client.get(f"/api/report/{report_id}", headers=headers)
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["score_breakdown"]
    assert data["enhanced_analysis"]
    assert "section_analysis" in data["enhanced_analysis"]


def test_job_match(client: TestClient, tmp_path: Path):
    headers = register_and_login(client)
    resume_id = upload_valid_resume(client, headers, tmp_path)
    response = client.post(
        "/api/analysis/job-match",
        headers=headers,
        json={
            "resume_id": resume_id,
            "job_title": "FastAPI Backend Developer",
            "job_description": "We need Python FastAPI SQLite Git Docker and strong communication skills.",
        },
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert "matched_keywords" in data
    assert "missing_keywords" in data
    assert "matched_skills" in data
    assert "missing_skills" in data
    assert "suggestions" in data


def test_report_history_and_pdf_download(client: TestClient, tmp_path: Path):
    headers = register_and_login(client)
    resume_id = upload_valid_resume(client, headers, tmp_path)
    analyze_response = client.post(f"/api/analysis/resume/{resume_id}", headers=headers)
    report_id = analyze_response.json()["data"]["report_id"]

    history_response = client.get("/api/report/history", headers=headers)
    assert history_response.status_code == 200
    history_data = history_response.json()["data"]
    assert len(history_data) == 1
    assert history_data[0]["id"] == report_id
    assert history_data[0]["report_id"] == report_id
    assert history_data[0]["resume_id"] == resume_id

    pdf_response = client.get(f"/api/report/{report_id}/download", headers=headers)
    assert pdf_response.status_code == 200
    assert pdf_response.headers["content-type"] == "application/pdf"


def test_admin_apis_require_admin_and_return_data(client: TestClient, tmp_path: Path):
    unauthenticated = client.get("/api/admin/overview")
    assert unauthenticated.status_code == 401

    user_headers = register_and_login(client)
    forbidden = client.get("/api/admin/overview", headers=user_headers)
    assert forbidden.status_code == 403

    resume_id = upload_valid_resume(client, user_headers, tmp_path)
    analyze_response = client.post(f"/api/analysis/resume/{resume_id}", headers=user_headers)
    assert analyze_response.status_code == 200

    admin_headers = create_admin_headers(client)
    overview_response = client.get("/api/admin/overview", headers=admin_headers)
    assert overview_response.status_code == 200
    overview = overview_response.json()["data"]
    assert overview["total_users"] == 2
    assert overview["total_resumes"] == 1
    assert overview["total_reports"] == 1
    assert overview["backend_status"] == "healthy"

    for path in ["/api/admin/users", "/api/admin/resumes", "/api/admin/reports", "/api/admin/logs"]:
        response = client.get(path, headers=admin_headers)
        assert response.status_code == 200
        assert response.json()["success"] is True

    report_id = analyze_response.json()["data"]["report_id"]
    report_response = client.get(f"/api/admin/reports/{report_id}", headers=admin_headers)
    assert report_response.status_code == 200
    assert report_response.json()["data"]["id"] == report_id
    assert report_response.json()["data"]["user"]["email"] == "test@example.com"

    forbidden_detail = client.get(f"/api/admin/reports/{report_id}", headers=user_headers)
    assert forbidden_detail.status_code == 403

    missing_response = client.get("/api/report/999999", headers=user_headers)
    assert missing_response.status_code == 404
    assert missing_response.json()["message"] == "Report not found or you do not have access to this report."


def test_activity_logs_are_created_for_core_actions(client: TestClient, tmp_path: Path):
    response = client.post(
        "/api/auth/register",
        json={"name": "Log User", "email": "log@example.com", "password": "Password123"},
    )
    assert response.status_code == 201

    failed_login = client.post("/api/auth/login", json={"email": "log@example.com", "password": "WrongPassword"})
    assert failed_login.status_code == 401

    login_response = client.post("/api/auth/login", json={"email": "log@example.com", "password": "Password123"})
    assert login_response.status_code == 200
    user_headers = {"Authorization": f"Bearer {login_response.json()['data']['access_token']}"}

    resume_id = upload_valid_resume(client, user_headers, tmp_path)
    analyze_response = client.post(f"/api/analysis/resume/{resume_id}", headers=user_headers)
    report_id = analyze_response.json()["data"]["report_id"]
    pdf_response = client.get(f"/api/report/{report_id}/download", headers=user_headers)
    assert pdf_response.status_code == 200
    logout_response = client.post("/api/auth/logout", headers=user_headers)
    assert logout_response.status_code == 200

    admin_headers = create_admin_headers(client)
    overview_response = client.get("/api/admin/overview", headers=admin_headers)
    assert overview_response.status_code == 200

    db = client.app.state.TestingSessionLocal()
    try:
        actions = {row.action for row in db.query(ActivityLog).all()}
    finally:
        db.close()

    assert {
        "user_register",
        "login_failed",
        "user_login",
        "resume_upload",
        "resume_analysis",
        "report_download",
        "logout",
        "admin_dashboard_access",
    }.issubset(actions)
