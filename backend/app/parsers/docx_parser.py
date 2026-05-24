from pathlib import Path

from docx import Document

from app.parsers.text_cleaner import clean_text


def extract_docx_text(file_path: str | Path) -> str:
    try:
        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                table_cells.extend(cell.text for cell in row.cells)
        text = "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise ValueError("Unable to parse DOCX resume. The file may be corrupted.") from exc
    cleaned = clean_text(text)
    if not cleaned:
        raise ValueError("DOCX resume text is empty after extraction.")
    return cleaned
